# Imports
#####################################################################
import docx
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import numpy as np
import pandas as pd

# PARAMETERS SECTION
#####################################################################
# sig_dig is set to dictate the precision for rounding numeric statistics
sig_dig = 1

# Assume df is the DataFrame you have
df = pd.read_csv('your_data_file.csv')  # You would set this to your actual DataFrame object here.

# Dictionary to define which columns are categorical (True) or not (False)
featured_categorical = {
    'gender': True,
    'main_discipline': True,
    #'migration_year': True,
    'destination_country': True,
    'preceding_country': True,
    'journal_counts': False,
    'publication_counts': False,
    'days_since_pub': False,
    'co-author_counts': False,
    'journal_counts': False,
    'country_count': False,
    'institute_count': False,
    'destination_coauthor_count': False,
    'destination_coauthor_num': False, 
    'author_subject_diversity': False,
    'author_coarse_subject_diversity': False,
    'sole_author_count': False,
}

# Headers for the table columns in the Word document
headers = ['Features', 'N (%)', 'Mean, Median, (Std. Dev.)\n[Range]']

#####################################################################
# Extract column names and their categorical status into lists
keep_col = list(featured_categorical.keys())
cat_col = list(featured_categorical.values())

# FORMAT FEATURE NAMES: Capitalizes each word in a camel-case or snake-case string
def formating(string):
    string =  ' '.join([word.capitalize() for word in string.split('_')])
    return  ' '.join([word.capitalize() for word in string.split(' ')])

# COMPUTE STATISTICS: Outputs stats for numerical features
def stating(feature):
    stats = feature.describe()  # Use pandas describe to get summary statistics
    out = str(np.round(stats.loc['mean'], sig_dig)) + ", "
    out += str(stats.loc['50%']) + ", "
    out += "(" + str(np.round(stats.loc['std'], sig_dig)) + "), "
    out += "\n[" + str(stats.loc['min']) + " - " + str(stats.loc['max']) + "]"
    return out

# Load an existing Word document template
doc = docx.Document('./RAND_template.docx')

line = 0  # Initialize line counter for table rows

# ADD HEADERS: Loop over header names and add them to the table
for h in range(len(headers)):
    header = headers[h]
    cell = doc.tables[0].cell(0, h)  # Access the cell at the table's first row

    # Clear existing paragraph and add a new one
    if cell.paragraphs:
        para = cell.paragraphs[0]
        para.clear()
    else:
        para = cell.add_paragraph()

    # Center-align subsequent headers (not the first column header)
    if h != 0:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add header text to the table cell with specified formatting
    run = para.add_run(header)
    run.font.size = Pt(9)
    run.font.name = 'Arial'
    run.bold = True

line += 1  # Move to the next line after adding headers

# ADD DATA TO THE TABLE
L = len(keep_col)  # Number of features to process
for i in range(L):
    col = keep_col[i]  # Current feature/column name
    cat = cat_col[i]  # Whether feature is categorical

    if cat == True:
        # CATEGORY LABEL: Add a row for the category label
        try:
            cell = doc.tables[0].cell(line, 0)
        except IndexError:
            doc.tables[0].add_row()  # Add a new row if index exceeds table's current rows
            cell = doc.tables[0].cell(line, 0)

        # Clear existing paragraph and add a new one
        if cell.paragraphs:
            para = cell.paragraphs[0]
            para.clear()
        else:
            para = cell.add_paragraph()

        # Set left indent and format for category
        para.paragraph_format.left_indent = Inches(0.0)
        run = para.add_run(formating(col))
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.bold = True
        line += 1

        # Calculate frequencies and percentages for each category item
        items, freqs = np.unique(df[col].astype(str), return_counts=True)
        percs = np.round(100 * (freqs / np.sum(freqs)), sig_dig)

        for k in range(len(items)):
            item = items[k]
            freq = freqs[k]
            perc = percs[k]

            # SUBCATEGORY TYPE: Add subcategory item for each unique value
            try:
                cell = doc.tables[0].cell(line, 0)
            except IndexError:
                doc.tables[0].add_row()
                cell = doc.tables[0].cell(line, 0)
            
            # Edit or add a new paragraph to the cell for subcategory
            if cell.paragraphs:
                para = cell.paragraphs[0]
                para.clear()
            else:
                para = cell.add_paragraph()

            para.paragraph_format.left_indent = Inches(0.25)
            run = para.add_run(str(item.capitalize()))
            run.font.size = Pt(9)
            run.font.name = 'Arial'

            # N COUNT: Frequency and percentage for the subcategory in the next column
            cell = doc.tables[0].cell(line, 1)
            if cell.paragraphs:
                para = cell.paragraphs[0]
                para.clear()
            else:
                para = cell.add_paragraph()

            para.paragraph_format.left_indent = Inches(0.0)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(f"{freq} ({perc}%)")
            run.font.size = Pt(9)
            run.font.name = 'Arial'

            line += 1
    
    elif cat == False:
        # FEATURE LABEL: Add a row for numerical features
        try:
            cell = doc.tables[0].cell(line, 0)  # Access cell to add feature label
        except IndexError:
            doc.tables[0].add_row()
            cell = doc.tables[0].cell(line, 0)

        # Clear existing paragraph and add new one for feature label
        if cell.paragraphs:
            para = cell.paragraphs[0]
            para.clear()
        else:
            para = cell.add_paragraph()

        para.paragraph_format.left_indent = Inches(0.0)
        run = para.add_run(formating(col))  # Format feature name
        run.font.size = Pt(9)
        run.font.name = 'Arial'
        run.bold = True

        # STATS VALUES: Add calculated stats to the next column
        stats = stating(df[col])

        cell = doc.tables[0].cell(line, 2)
        if cell.paragraphs:
            para = cell.paragraphs[0]
            para.clear()
        else:
            para = cell.add_paragraph()

        para.paragraph_format.left_indent = Inches(0.0)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(stats)  # Display statistics result
        run.font.size = Pt(9)
        run.font.name = 'Arial'

        line += 1

# Save the document after processing
doc.save('./RAND_summary_statistics.docx')